import os
import time
import tempfile
from yt_dlp import YoutubeDL
from transformers import pipeline
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document
from docx.shared import Pt

# Step 1: Extract Audio from YouTube Video
def download_audio(youtube_url, temp_dir):
    try:
        ydl_opts = {
            'format': 'bestaudio/best',  # Download the best quality audio
            'outtmpl': os.path.join(temp_dir, 'audio.%(ext)s'),  # Output file name in temp directory
            'postprocessors': [{
                'key': 'FFmpegExtractAudio',  # Extract audio using FFmpeg
                'preferredcodec': 'mp3',     # Convert to MP3
                'preferredquality': '192',   # Audio quality
            }],
            'quiet': True,  # Suppress yt-dlp output
            'ffmpeg_location': 'C:/ffmpeg/ffmpeg.exe'  # Specify FFmpeg path (Windows example)
        }
        with YoutubeDL(ydl_opts) as ydl:
            ydl.download([youtube_url])
        audio_file = os.path.join(temp_dir, 'audio.mp3')
        return audio_file  # Return the downloaded audio file name
    except Exception as e:
        print(f"Error downloading audio: {e}")
        return None

# Step 2: Transcribe Audio to Text
def transcribe_audio(audio_file):
    try:
        # Use the Hugging Face ASR pipeline with return_timestamps enabled
        transcriber = pipeline(
            "automatic-speech-recognition",
            model="openai/whisper-small",
            return_timestamps=True,  # Enable timestamps for long-form audio
            # generate_kwargs={"language": "english"}  # Always translate to English
        )
        transcription = transcriber(audio_file)
        return transcription['text']
    except Exception as e:
        print(f"Error transcribing audio: {e}")
        return None

# Step 3: Summarize the Text
def summarize_text(text):
    try:
        # Use a summarization model (e.g., microsoft/Phi-3-mini-4k-instruct or facebook/bart-large-cnn)
        summarizer = pipeline("summarization", model="facebook/bart-large-cnn")
        summary = summarizer(text, max_length=430, min_length=30, do_sample=False)
        return summary[0]['summary_text']
    except Exception as e:
        print(f"Error summarizing text: {e}")
        return None

# Step 4: Generate Professional PDF
def generate_pdf(summary, transcript, filename="summary.pdf"):
    try:
        # Create a PDF document
        doc = SimpleDocTemplate(filename, pagesize=letter)
        styles = getSampleStyleSheet()
        content = []

        # Add title
        title = Paragraph("Video Summary", styles['Title'])
        content.append(title)
        content.append(Spacer(1, 12))

        # Add summary
        summary_para = Paragraph(summary, styles['BodyText'])
        content.append(summary_para)
        content.append(Spacer(1, 12))

        # Add transcript heading
        transcript_title = Paragraph("Transcript", styles['Heading2'])
        content.append(transcript_title)
        content.append(Spacer(1, 12))

        # Add transcript
        transcript_para = Paragraph(transcript, styles['BodyText'])
        content.append(transcript_para)

        # Build the PDF
        doc.build(content)
        print(f"PDF generated: {filename}")
    except Exception as e:
        print(f"Error generating PDF: {e}")

# Step 5: Generate Professional DOCX
def generate_docx(summary, transcript, filename="summary.docx"):
    try:
        # Create a Word document
        doc = Document()

        # Add title
        doc.add_heading('Video Summary', 0)

        # Add summary
        doc.add_heading('Summary', level=1)
        doc.add_paragraph(summary)

        # Add transcript
        doc.add_heading('Transcript', level=1)
        doc.add_paragraph(transcript)

        # Save the document
        doc.save(filename)
        print(f"DOCX generated: {filename}")
    except Exception as e:
        print(f"Error generating DOCX: {e}")

# Retry mechanism for file deletion
def delete_file_with_retry(file_path, retries=5, delay=2):
    for i in range(retries):
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
                print(f"Deleted file: {file_path}")
                return True
        except Exception as e:
            print(f"Attempt {i + 1}: Error deleting file {file_path} - {e}")
            time.sleep(delay)
    return False

# Main Function
def main(youtube_url):
    # Create a temporary directory for the audio file
    with tempfile.TemporaryDirectory() as temp_dir:
        # Step 1: Download audio
        print("Downloading the audio")
        audio_file = download_audio(youtube_url, temp_dir)
        if not audio_file:
            return

        print("Transcribing the audio")
        # Step 2: Transcribe audio to text
        transcript = transcribe_audio(audio_file)
        if not transcript:
            return

        # Step 3: Summarize the text
        print("Summarizing the text")
        summary = summarize_text(transcript)
        if not summary:
            return

        print("Converting into pdf")
        # Step 4: Generate PDF and DOCX documents
        generate_pdf(summary, transcript, "video_summary_.pdf")
        print("Converting into docx")
        generate_docx(summary, transcript, "video_summary_.docx")

        # Clean up
        if not delete_file_with_retry(audio_file):
            print(f"Failed to delete audio file: {audio_file}")

if __name__ == "__main__":
    youtube_url = input("Enter the YouTube video URL: ")
    print("Started working on it plz wait !!!")
    main(youtube_url)
